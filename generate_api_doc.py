from docx import Document

content = {
    'title': 'Amazon SP-API Emulator API Reference',
    'base_url': 'http://localhost:4000',
    'endpoints': [
        {
            'title': '1) Authentication',
            'method': 'POST',
            'url': 'http://localhost:4000/auth/token',
            'headers': 'None',
            'body': 'None',
            'response': '{\n  "access_token": "fake-token",\n  "refresh_token": "xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx",\n  "expires_in": 3600,\n  "token_type": "Bearer"\n}'
        },
        {
            'title': '2) Create / Submit a Listing',
            'method': 'PUT',
            'url': 'http://localhost:4000/listings/2021-08-01/items/:sellerId/:sku',
            'headers': 'Authorization: Bearer <access_token>\nContent-Type: application/json',
            'body': '{\n  "productType": "LAPTOP",\n  "attributes": {\n    "brand": [{ "value": "Dell" }],\n    "item_name": [{ "value": "XPS 13 Laptop" }]\n  },\n  "webhookUrl": "https://webhook.site/YOUR_REAL_ID"\n}',
            'response': '{\n  "sku": "XPS13",\n  "submissionId": "...",\n  "status": "ACCEPTED"\n}'
        },
        {
            'title': '3) List All Listings',
            'method': 'GET',
            'url': 'http://localhost:4000/listings',
            'headers': 'Authorization: Bearer <access_token>',
            'body': 'None'
        },
        {
            'title': '4) Get Listing by SKU',
            'method': 'GET',
            'url': 'http://localhost:4000/listings/:sku',
            'example': 'http://localhost:4000/listings/XPS13',
            'headers': 'Authorization: Bearer <access_token>',
            'body': 'None'
        },
        {
            'title': '5) Delete Listing by ID',
            'method': 'DELETE',
            'url': 'http://localhost:4000/listings/id/:id',
            'headers': 'Authorization: Bearer <access_token>',
            'body': 'None'
        },
        {
            'title': '6) Delete Listing by SKU',
            'method': 'DELETE',
            'url': 'http://localhost:4000/listings/:sku',
            'example': 'http://localhost:4000/listings/XPS13',
            'headers': 'Authorization: Bearer <access_token>',
            'body': 'None'
        },
        {
            'title': '7) Update Inventory Quantity',
            'method': 'PATCH',
            'url': 'http://localhost:4000/inventory/:sku',
            'example': 'http://localhost:4000/inventory/XPS13',
            'headers': 'Authorization: Bearer <access_token>\nContent-Type: application/json',
            'body': '{\n  "quantity": 10\n}'
        },
        {
            'title': '8) Get Inventory Item by SKU',
            'method': 'GET',
            'url': 'http://localhost:4000/inventory/:sku',
            'example': 'http://localhost:4000/inventory/XPS13',
            'headers': 'Authorization: Bearer <access_token>',
            'body': 'None'
        }
    ],
    'dashboard': [
        'http://localhost:4000/dashboard/analytics',
        'http://localhost:4000/dashboard/events',
        'http://localhost:4000/dashboard/processing-status',
        'http://localhost:4000/dashboard/queue-metrics',
        'http://localhost:4000/dashboard/webhooks'
    ],
    'notes': [
        'Only /auth/token is public.',
        'All other API routes require Authorization: Bearer <access_token>.',
        'Use a real webhook.site URL for webhookUrl, not placeholder text.',
        'Example valid webhookUrl: https://webhook.site/<real-generated-id>'
    ]
}

doc = Document()
doc.add_heading(content['title'], level=1)
doc.add_paragraph(f"Base URL: {content['base_url']}")

for endpoint in content['endpoints']:
    doc.add_heading(endpoint['title'], level=2)
    doc.add_paragraph(f"Method: {endpoint['method']}")
    doc.add_paragraph(f"URL: {endpoint['url']}")
    if 'example' in endpoint:
        doc.add_paragraph(f"Example: {endpoint['example']}")
    doc.add_paragraph(f"Headers: {endpoint['headers']}")
    doc.add_paragraph(f"Body: {endpoint['body']}")
    if 'response' in endpoint:
        doc.add_paragraph('Response example:')
        doc.add_paragraph(endpoint['response'])

    doc.add_paragraph('')

if content['dashboard']:
    doc.add_heading('Dashboard / Monitoring Endpoints', level=2)
    doc.add_paragraph('All dashboard endpoints are GET and require Authorization: Bearer <access_token>.')
    for line in content['dashboard']:
        doc.add_paragraph(line)

if content['notes']:
    doc.add_heading('Notes', level=2)
    for note in content['notes']:
        doc.add_paragraph(note, style='List Bullet')

filename = 'c:\\Users\\Admin\\Downloads\\amazon-sandbox\\api-endpoints.docx'
doc.save(filename)
print(f'WROTE {filename}')
